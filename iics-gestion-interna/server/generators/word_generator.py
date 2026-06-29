"""
Generador de Documentos Word con Formato APA 7
Sistema integrado con el Asistente HOYR
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os
from pathlib import Path


class WordGenerator:
    """Generador profesional de documentos Word con formato APA 7"""
    
    def __init__(self, logos_path='public/certificates'):
        self.logos_path = Path(logos_path)
        self.doc = Document()
        self._configurar_formato_apa7()
    
    def _configurar_formato_apa7(self):
        """Configura el documento con formato APA 7"""
        # Márgenes: 1 pulgada (2.54 cm) en todos los lados
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Fuente predeterminada: Times New Roman 12pt
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Espaciado de línea: doble
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    def agregar_encabezado_logos(self):
        """Agrega logos institucionales al encabezado"""
        section = self.doc.sections[0]
        header = section.header
        
        # Tabla 1x3 para logos (izq, centro, derecha)
        table = header.add_table(1, 3, width=Inches(6.5))
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Logo UNC (izquierda)
        logo_unc = self.logos_path / 'logo-unc' / 'R.png'
        if logo_unc.exists():
            cell_left = table.rows[0].cells[0]
            para = cell_left.paragraphs[0]
            run = para.add_run()
            run.add_picture(str(logo_unc), width=Inches(0.8))
        
        # Texto central
        cell_center = table.rows[0].cells[1]
        para = cell_center.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run('UNIVERSIDAD NACIONAL DE CAJAMARCA\nFacultad de Ciencias Sociales')
        run.font.size = Pt(10)
        run.font.bold = True
        
        # Logo Facultad (derecha)
        logo_facultad = self.logos_path / 'logo-facultad' / 'logo-facultad.png'
        if logo_facultad.exists():
            cell_right = table.rows[0].cells[2]
            para = cell_right.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = para.add_run()
            run.add_picture(str(logo_facultad), width=Inches(0.8))
    
    def agregar_pie_pagina(self):
        """Agrega logo de revista y número de página al pie"""
        section = self.doc.sections[0]
        footer = section.footer
        
        # Logo Revista ACS
        logo_revista = self.logos_path / 'logo-revista' / 'logo-revista-ACS.png'
        if logo_revista.exists():
            para = footer.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(str(logo_revista), height=Inches(0.5))
        
        # Número de página
        para = footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        # Código de campo para número de página
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
    
    def agregar_portada(self, titulo, autores, fecha=None):
        """Crea portada estilo APA 7"""
        if fecha is None:
            fecha = datetime.now().strftime('%d de %B de %Y')
        
        # Título
        para_titulo = self.doc.add_paragraph()
        para_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para_titulo.add_run(titulo)
        run.font.size = Pt(16)
        run.font.bold = True
        
        # Espaciado
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Autores
        para_autores = self.doc.add_paragraph()
        para_autores.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if isinstance(autores, list):
            autores_texto = '\n'.join(autores)
        else:
            autores_texto = autores
        run = para_autores.add_run(autores_texto)
        run.font.size = Pt(12)
        
        # Espaciado
        self.doc.add_paragraph()
        
        # Institución
        para_inst = self.doc.add_paragraph()
        para_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para_inst.add_run('Universidad Nacional de Cajamarca\nFacultad de Ciencias Sociales')
        run.font.size = Pt(12)
        
        # Fecha
        self.doc.add_paragraph()
        para_fecha = self.doc.add_paragraph()
        para_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para_fecha.add_run(fecha)
        run.font.size = Pt(12)
        
        # Salto de página
        self.doc.add_page_break()
    
    def agregar_seccion(self, titulo, nivel=1):
        """Agrega un encabezado de sección"""
        if nivel == 1:
            para = self.doc.add_heading(titulo, level=1)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            para = self.doc.add_heading(titulo, level=nivel)
        
        return para
    
    def agregar_parrafo(self, texto, negrita=False, cursiva=False):
        """Agrega un párrafo con formato"""
        para = self.doc.add_paragraph()
        run = para.add_run(texto)
        
        if negrita:
            run.font.bold = True
        if cursiva:
            run.font.italic = True
        
        return para
    
    def agregar_tabla(self, datos, encabezados=None, colores=True):
        """
        Agrega tabla profesional
        datos: lista de listas [[fila1], [fila2], ...]
        encabezados: lista de strings ['Col1', 'Col2', ...]
        """
        rows = len(datos) + (1 if encabezados else 0)
        cols = len(datos[0]) if datos else 0
        
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Light Grid Accent 1'
        
        # Encabezados
        if encabezados:
            for i, encabezado in enumerate(encabezados):
                cell = table.rows[0].cells[i]
                cell.text = encabezado
                # Color de fondo azul UNC
                if colores:
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), '003366')  # Azul oscuro UNC
                    cell._element.get_or_add_tcPr().append(shading_elm)
                    # Texto blanco
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.bold = True
        
        # Datos
        start_row = 1 if encabezados else 0
        for i, fila in enumerate(datos):
            for j, valor in enumerate(fila):
                table.rows[start_row + i].cells[j].text = str(valor)
        
        return table
    
    def guardar(self, nombre_archivo):
        """Guarda el documento"""
        filepath = Path(nombre_archivo)
        filepath.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(filepath))
        return str(filepath)


# Ejemplo de uso para testing
if __name__ == '__main__':
    gen = WordGenerator()
    gen.agregar_encabezado_logos()
    gen.agregar_pie_pagina()
    gen.agregar_portada(
        titulo='INFORME EJECUTIVO DE PRUEBA',
        autores=['Edwar Jhanpiere Saenz Tello'],
        fecha='28 de enero de 2026'
    )
    gen.agregar_seccion('Introducción', nivel=1)
    gen.agregar_parrafo('Este es un documento de prueba generado automáticamente.')
    gen.agregar_seccion('Datos', nivel=1)
    gen.agregar_tabla(
        datos=[
            ['Juan', '85', 'Aprobado'],
            ['María', '92', 'Aprobado'],
            ['Pedro', '78', 'Aprobado']
        ],
        encabezados=['Estudiante', 'Nota', 'Estado']
    )
    
    gen.guardar('test_documento.docx')
    print("[OK] Documento de prueba generado: test_documento.docx")

